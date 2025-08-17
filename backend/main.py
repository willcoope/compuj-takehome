import json
import io
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import create_engine, Column, Integer, String, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from datetime import datetime
from transformers import pipeline
import PyPDF2
from docx import Document as DocxDocument
from pydantic import BaseModel, RootModel
from typing import Dict, Any, List
from sqlalchemy.orm import Session # Import Session for type hinting
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Constants for chunking
MAX_CHUNK_SIZE = 1000  # characters, roughly corresponds to model's token limit (1024 tokens)
CHUNK_OVERLAP = 100    # characters, for context preservation between chunks
LOW_CONFIDENCE_THRESHOLD = 0.45 # Threshold for flagging low confidence classifications

# Helper function for text chunking
def chunk_text(text: str) -> List[str]:
    chunks = []
    if len(text) <= MAX_CHUNK_SIZE:
        return [text]

    start = 0
    while start < len(text):
        end = start + MAX_CHUNK_SIZE
        current_chunk = text[start:end]
        chunks.append(current_chunk)
        start += MAX_CHUNK_SIZE - CHUNK_OVERLAP
    return chunks

# Zero-shot classification model
logger.info("Loading zero-shot classification model...")
classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")
logger.info("Model loaded successfully.")
candidate_labels = [
    "Technical Documentation",
    "Business Proposal",
    "Legal Document",
    "Academic Paper",
    "General Article",
    "Other"
]

# Database configuration
SQLALCHEMY_DATABASE_URL = "sqlite:///./documents.db"
engine = create_engine(
    SQLALCHEMY_DATABASE_URL, connect_args={"check_same_thread": False}
)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# Database model
class Document(Base):
    __tablename__ = "documents"

    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String, index=True)
    content = Column(String)
    upload_time = Column(DateTime, default=datetime.now)
    predicted_category = Column(String, nullable=True)
    confidence_scores = Column(String, nullable=True) # Stored as JSON string

# Dependency to get DB session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# Create database tables only if the script is run directly
# This prevents table creation during tests when main.py is imported
if __name__ == "__main__":
    logger.info("Creating database tables...")
    Base.metadata.create_all(bind=engine)
    logger.info("Database tables created.")

app = FastAPI()

origins = [
    "http://localhost:3000",  # React app default port
    "http://127.0.0.1:3000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/health")
async def health_check():
    logger.info("Health check requested.")
    return {"status": "ok"}

# Pydantic Models
class ConfidenceScores(RootModel[Dict[str, float]]):
    pass

class DocumentResponse(BaseModel):
    id: int
    filename: str
    upload_time: datetime
    predicted_category: str
    confidence_scores: ConfidenceScores

class UploadResponse(BaseModel):
    message: str
    filename: str
    upload_time: datetime
    predicted_category: str
    confidence_scores: ConfidenceScores

@app.post("/upload", response_model=UploadResponse)
async def upload_document(file: UploadFile = File(...), db: Session = Depends(get_db)):
    logger.info(f"Received upload request for file: {file.filename}")
    content_str = ""
    filename = file.filename

    if filename.endswith(".txt"):
        content = await file.read()
        content_str = content.decode("utf-8")
        logger.info(f"Extracted text from .txt file: {filename}")
    elif filename.endswith(".pdf"):
        try:
            # Read PDF content
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(await file.read()))
            for page_num in range(len(pdf_reader.pages)):
                content_str += pdf_reader.pages[page_num].extract_text() or ""
            logger.info(f"Extracted text from .pdf file: {filename}")
        except Exception as e:
            logger.error(f"Error processing PDF file {filename}: {e}")
            raise HTTPException(status_code=400, detail=f"Error processing PDF: {e}")
    elif filename.endswith(".docx"):
        try:
            # Read DOCX content
            doc = DocxDocument(io.BytesIO(await file.read()))
            for paragraph in doc.paragraphs:
                content_str += paragraph.text + "\n"
            logger.info(f"Extracted text from .docx file: {filename}")
        except Exception as e:
            logger.error(f"Error processing DOCX file {filename}: {e}")
            raise HTTPException(status_code=400, detail=f"Error processing DOCX: {e}")
    else:
        logger.warning(f"Unsupported file type uploaded: {filename}")
        raise HTTPException(status_code=400, detail="Only .txt, .pdf, and .docx files are allowed")

    # --- ML Pipeline with Chunking and Aggregation ---
    logger.info(f"Classifying document {filename}...")
    # Chunk the text
    text_chunks = chunk_text(content_str)
    
    aggregated_scores = {label: 0.0 for label in candidate_labels}
    num_chunks = len(text_chunks)

    for i, chunk in enumerate(text_chunks):
        if chunk.strip(): # Ensure chunk is not empty or just whitespace
            try:
                classification_results = classifier(chunk, candidate_labels)
                for label, score in zip(classification_results['labels'], classification_results['scores']):
                    aggregated_scores[label] += score
                logger.debug(f"Chunk {i+1}/{num_chunks} classified.")
            except Exception as e:
                logger.error(f"Error classifying chunk {i+1}/{num_chunks} for {filename}: {e}")
                # Decide how to handle classification errors per chunk - currently continues
    
    # Average the scores across chunks
    if num_chunks > 0: # Avoid division by zero if all chunks were empty
        for label in aggregated_scores:
            aggregated_scores[label] /= num_chunks
    else:
        logger.warning(f"No valid chunks to classify for {filename}. Setting predicted category to 'Other'.")

    # Sort aggregated scores to find the predicted category and prepare for storage
    sorted_aggregated_scores = sorted(aggregated_scores.items(), key=lambda item: item[1], reverse=True)
    
    # Determine predicted category and confidence scores for storage
    predicted_category = sorted_aggregated_scores[0][0] if sorted_aggregated_scores else "Other"
    confidence_scores = {label: score for label, score in sorted_aggregated_scores}

    # --- Low Confidence Handling ---
    top_confidence = confidence_scores.get(predicted_category, 0.0)
    if top_confidence < LOW_CONFIDENCE_THRESHOLD and predicted_category != "Other":
        logger.info(f"Low confidence classification for {filename}: {predicted_category} ({top_confidence:.2f}). Reclassifying to 'Other'.")
        predicted_category = "Other"
    logger.info(f"Document {filename} classified as: {predicted_category} (Confidence: {top_confidence:.2f})")

    try:
        db_document = Document(
            filename=filename,
            content=content_str,
            predicted_category=predicted_category,
            confidence_scores=json.dumps(confidence_scores)
        )
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        logger.info(f"Document {filename} saved to database with ID: {db_document.id}")
        return {
            "message": "File uploaded successfully",
            "filename": db_document.filename,
            "upload_time": db_document.upload_time,
            "predicted_category": db_document.predicted_category,
            "confidence_scores": json.loads(db_document.confidence_scores)
        }
    except Exception as e:
        db.rollback() # Rollback changes on error
        logger.error(f"Database error when saving document {filename}: {e}")
        raise HTTPException(status_code=500, detail="Database error when saving document.")
    finally:
        db.close()

@app.get("/documents", response_model=List[DocumentResponse])
async def get_documents(db: Session = Depends(get_db)):
    logger.info("Fetching all documents from database.")
    try:
        documents = db.query(Document).all()
        logger.info(f"Retrieved {len(documents)} documents.")
        return [
            {
                "id": doc.id,
                "filename": doc.filename,
                "upload_time": doc.upload_time,
                "predicted_category": doc.predicted_category,
                "confidence_scores": json.loads(doc.confidence_scores)
            }
            for doc in documents
        ]
    except Exception as e:
        logger.error(f"Database error when fetching documents: {e}")
        raise HTTPException(status_code=500, detail="Database error when fetching documents.")
    finally:
        db.close()
